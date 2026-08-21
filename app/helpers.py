import re
import json
from datetime import datetime

from flask import g, current_app
from sqlalchemy import (
    select,
)
from docx import Document
from docx.shared import Pt, Mm, Cm, Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
import yaml

from bs4 import BeautifulSoup
import requests

import boto3
from botocore.exceptions import ClientError

from app.models import (
    Collection,
    Publication,
    PublicationLiterature,
    Item,
    ItemSynonym,
)
from app.database import session
from app.jinja_func import pick_first, italicize_name

TAIWAN_COUNTIES = {
    '宜蘭縣': 'Yilan',
    '桃園市': 'Taoyuan',
    '新竹縣': 'Hsinchu',
    '苗栗縣': 'Miaoli',
    '彰化縣': 'Changhua',
    '南投縣': 'Nantou',
    '雲林縣': 'Yunlin',
    '嘉義縣': 'Chiayi',
    '屏東縣': 'Pingtung',
    '臺東縣': 'Taitung',
    '花蓮縣': 'Hualien',
    '澎湖縣': 'Penghu',
    '基隆市': 'Keelung',
    '新竹市': 'Hsinchu',
    '嘉義市': 'Chiayi',
    '臺北市': 'Taipei',
    '新北市': 'New Taipei',
    '臺中市': 'Taichung',
    '臺南市': 'Tainan',
    '高雄市': 'Kaohsiung'
}

class BiotaPrint(object):
    doc = None

    def __init__(self):
        self.doc = Document()

    def as_docx(self):
        return self.doc

    def save(self, name):
        self.doc.save(f'{name}.docx')

    def create_column_section(self, num_columns):
        """Add section with specified number of columns and optional custom widths."""
        section = self.doc.add_section(WD_SECTION_START.CONTINUOUS)

        page_width = 8.5  # Standard US Letter width in inches
        margin = 1.0      # 1-inch margins
        usable_width = page_width - 2 * margin  # Available width for content
        column_widths = [usable_width] if num_columns == 1 else [usable_width / 2, usable_width / 2]

        # Create the columns XML element
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), str(num_columns))
        cols.set(qn('w:equalWidth'), "0" if column_widths else "1")

        # If specific widths are provided, add them to the columns element
        if column_widths:
            for width in column_widths:
                col = OxmlElement('w:col')
                # Convert inches to twentieths of a point (unit used in docx)
                width_in_twips = int(width * 1440)  # 1440 twips = 1 inch
                col.set(qn('w:w'), str(width_in_twips))
                cols.append(col)

        # Add the columns element to the section properties
        section._sectPr.append(cols)
        return section

    def parse_styled_text(self, text):

        result = []
        text = text.replace('<br>', '\n\n')
        parts = re.split(r'(</?i>)', text)

        is_italic = False
        for part in parts:
            if part == '<i>':
                is_italic = True
            elif part == '</i>':
                is_italic = False
            elif part:  # Keep the text segment as-is
                style = 'italic' if is_italic else 'normal'
                result.append([part, style])

        return result


    def add_box(self, data, args={}):

        is_list = False
        if isinstance(data, list):
            is_list = True

        p = self.doc.add_paragraph()

        align = args.get('align', '')
        styles = args.get('styles', [])
        size = args.get('size', '')
        custom = args.get('custom', '')

        if align == 'justify':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        runs = []

        if is_list:
            for x in data:
                res = self.parse_styled_text(x)
                runs += res
                runs.append(['\n', 'normal'])

        else:
            res = self.parse_styled_text(data)
            runs = res

        for t in runs:
            run = p.add_run(t[0])

            if 'h' in size:
                if 'bold' not in styles:
                    styles.append('bold')

            # auto add serial number (scientific name) to bold
            sn_pattern = r"^[0-9]+\. $"
            if m := re.match(sn_pattern, t[0]):
                run.bold = True

            if 'italic' in styles or t[1] == 'italic':
                run.italic = True

                if custom and custom == 'italic-to-bold':
                    run.italic = False
                    run.bold = True

            if 'bold' in styles:
                run.bold = True

            if size == 'h2':
                run.font.size = Pt(12)
            elif size == 'h3':
                run.font.size = Pt(11)
            else:
                if size:
                    run.font.size = Pt(int(size))


    def add_list(self, data, has_italic=False):
        p = self.doc.add_paragraph()
        for x in data:
            self.add_content(x + '\n', align='left', paragraph=p, has_italic=has_italic)


def generate_docx(data):

    biota = BiotaPrint()

    # default section
    biota.add_box('Biota Taiwanica', {'size':'h3'})
    biota.add_box(f'generated: {datetime.now()}')
    biota.doc.add_page_break()

    for idx, d in enumerate(data):
        section = biota.create_column_section(1)
        biota.add_box(d['title'], {'size': 'h2', 'align': 'center'})
        biota.add_box(d['author'], {'size': 'h3', 'align': 'center'})
        biota.add_box('LITERATURE', {'size': 'h3'})
        biota.add_box(d['literatures'])

        #biota.add_content(['a <i>haha</i> b', 'foo<i>o</i>, and <i>oo</i>p'])

        section = biota.create_column_section(2)

        counter = 0
        for i, v in enumerate(d['items']):
            if str(v['rank_id']) == '34': # only species has sort number
                counter += 1
                title = f"{counter}. {v['scientificName']}"
            else:
                title = v['scientificName']

            biota.add_box(title, {'size': '11', 'custom': 'italic-to-bold'})

            if v['commonNames']:
                x = pick_first(v['commonNames'], '|', 'zh')
                biota.add_box(x)

            #biota.add_content('SYNONYMS', 'h3')
            for x in v['synonyms']:
                biota.add_box(x)

            if x := v.get('description'):
                #biota.add_content('DESCRIPTION', 'h3')
                biota.add_box(x, {'align': 'justify'})

            if x:= v.get('distribution'):
                #biota.add_content('DISTRIBUTION', 'h3')
                biota.add_box(x, {'align': 'justify'})

            if len(v['specimens']):
                #biota.add_content('SPECIMENS', 'h3')
                #    biota.add_box(v['specimens'])
                s = ''
                for dist, sp_list in v['specimens'].items():
                    #s = dist
                    sp_arr = [x[1] for x in sp_list]
                    sp_str = ';'.join(sp_arr)
                    s += f'{dist}: {sp_str}. '
                biota.add_box(s)

            if x:= v['note']:
                #biota.add_content('NOTE', 'h3')
                biota.add_box(x, {'align': 'justify'})

    return biota.as_docx()


def generate_json(data):
    """Build a structured, presentation-free context from namespace data.

    This is the single source of truth for *what* gets published; renderers
    (PDF, JSON download) only decide *how* it looks. Everything here is plain
    JSON-serializable data: no fonts, styles or layout.

    Inline markup (<i>, <b>, <br>) is kept as-is inside text fields, since it
    carries taxonomic meaning rather than styling.
    """
    def _strip_italic(text):
        return (text or '').replace('<i>', '').replace('</i>', '').strip()

    publications = []

    for d in data:
        pub = {
            'title': d.get('title', ''),
            'author': d.get('author', ''),
            'category': None,
            'literatures': [],
            'keys': [],
            'items': [],
        }

        # category (higher rank than species), shown as publication heading
        if cat := d.get('item_category'):
            c = cat[0]
            # HACK: clean name, (source_data.name_authors)
            names = c.get('scientificName', '').split(',')
            scientific_name = _strip_italic(names[0])
            common_names = c.get('commonNames', '') or ''
            authors = ''
            if sd := c.get('source_data'):
                authors = sd.get('name_authors', '') or ''

            # heading: {scientificName} {authors} {commonNames}, skipping the
            # parts that are empty
            heading = ' '.join(
                x for x in (scientific_name, authors, common_names) if x)

            pub['category'] = {
                'scientificName': scientific_name,
                'commonNames': common_names,
                'authors': authors,
                'heading': heading,
                'description': c.get('description', '') or '',
            }

        # literatures: normalize dict/str to plain text
        for lit in d.get('literatures', []):
            content = lit.get('content', '') if isinstance(lit, dict) else str(lit)
            if content:
                pub['literatures'].append(content)

        # identification keys
        for key in d.get('keys', []):
            key_data = {
                'title': key.get('title', ''),
                'entries': [],
            }
            for entry in key.get('entries', []):
                # result: species name takes precedence over next couplet
                result = ''
                result_type = ''
                if name := entry.get('result_item_name'):
                    result = name
                    result_type = 'item'
                elif couplet := entry.get('result_couplet'):
                    result = couplet
                    result_type = 'couplet'

                key_data['entries'].append({
                    'number': entry.get('number', ''),
                    'indentLevel': entry.get('indent_level', 0),
                    'description': entry.get('description', '') or '',
                    'result': result,
                    'resultType': result_type,
                })
            pub['keys'].append(key_data)

        # items (species)
        counter = 0
        for v in d.get('items', []):
            counter += 1

            # HACK: split name, the part after the canonical name (author, ref...)
            # split on the LAST </i>: infraspecific names italicize more than
            # one part (<i>Camphora officinarum</i> var. <i>nominale</i> ...),
            # and splitting on the first one would cut the name in half and
            # leave an unbalanced <i> for ReportLab to choke on
            name_suffix = ''
            if full_name := v.get('fullScientificName'):
                parts = full_name.rsplit('</i>', 1)
                if len(parts) > 1:
                    name_suffix = parts[1].strip()

            common_name = ''
            if x := v.get('commonNames'):
                common_name = pick_first(x, '|', 'zh') or ''

            specimens = []
            for county, sp_list in (v.get('specimens') or {}).items():
                records = [x[1] for x in sp_list]
                if not records:
                    continue
                specimens.append({
                    'county': county,
                    'region': TAIWAN_COUNTIES.get(county, 'Other').upper(),
                    'records': records,
                })

            pub['items'].append({
                'number': counter,
                'scientificName': v.get('scientificName', ''),
                'nameSuffix': name_suffix,
                'fullScientificName': v.get('fullScientificName', ''),
                'rankId': v.get('rank_id', ''),
                'commonName': common_name,
                'synonyms': [x for x in v.get('synonyms', []) if x],
                'description': v.get('description', '') or '',
                'distribution': v.get('distribution', '') or '',
                'specimens': specimens,
                'note': v.get('note', '') or '',
            })

        publications.append(pub)

    return {
        'generator': 'Biota Taiwanica',
        'generatedAt': datetime.now().strftime('%Y-%m-%d %H:%M'),
        'publications': publications,
    }


def sanitize_html(text):
    """Sanitize HTML content for ReportLab compatibility."""
    if not text:
        return ''
    # Convert br tags to self-closing
    text = re.sub(r'<br\s*>', '<br/>', text, flags=re.IGNORECASE)
    # Escape special characters but preserve valid HTML tags
    # ReportLab supports: b, i, u, strike, super, sub, br, a
    return text


def convert_html_to_custom_fonts(text, base_font='Tinos'):
    """Convert HTML tags to font tags for custom font support.

    Args:
        text: HTML text with <b> and <i> tags
        base_font: Base font family ('Tinos' or 'Serif' for NotoSerifTC)

    Returns:
        Text with <font> tags instead of <b>/<i>
    """
    if not text:
        return ''

    fonts_map = {
        'regular': 'Serif-Regular',
        'bold': 'Serif-Bold',
        'italic': 'Serif-Italic',  # No italic variant for NotoSerifTC
        'bold-italic': 'Serif-Bold'
    }

    # Handle nested <b><i>...</i></b> or <i><b>...</b></i> -> bold-italic
    text = re.sub(
        r'<b>\s*<i>(.*?)</i>\s*</b>|<i>\s*<b>(.*?)</b>\s*</i>',
        lambda m: f'<font name="{fonts_map["bold-italic"]}">{m.group(1) or m.group(2)}</font>',
        text,
        flags=re.DOTALL
    )

    # Handle remaining <b>...</b> -> bold
    text = re.sub(
        r'<b>(.*?)</b>',
        lambda m: f'<font name="{fonts_map["bold"]}">{m.group(1)}</font>',
        text,
        flags=re.DOTALL
    )

    # Handle remaining <i>...</i> -> italic
    text = re.sub(
        r'<i>(.*?)</i>',
        lambda m: f'<font name="{fonts_map["italic"]}">{m.group(1)}</font>',
        text,
        flags=re.DOTALL
    )

    return text


def register_pdf_fonts():
    """Register the bundled TTF fonts used by the PDF renderers."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    custom_fonts = {
        'regular_tc': ['Serif-Regular', 'app/fonts/NotoSerifTC-Light.ttf'],
        # 400-weight TC face: matches NotoSerif-Italic, which <i> maps to, so a
        # mixed italic/upright line does not look two-toned
        'text_tc': ['Serif-Text', 'app/fonts/NotoSerifTC-Regular.ttf'],
        # 500-weight TC face, for emphasis that should stay short of bold
        'medium_tc': ['Serif-Medium', 'app/fonts/NotoSerifTC-Medium.ttf'],
        'bold_tc': ['Serif-Bold', 'app/fonts/NotoSerifTC-Bold.ttf'],
        'regular': ['Tinos-Regular', 'app/fonts/NotoSerif-Regular.ttf'],
        'bold': ['Tinos-Bold', 'app/fonts/Tinos-Bold.ttf'],
        'italic': ['Serif-Italic', 'app/fonts/NotoSerif-Italic.ttf'],
        'bold-italic': ['Tinos-BoldItalic', 'app/fonts/Tinos-BoldItalic.ttf'],
    }
    for v in custom_fonts.values():
        pdfmetrics.registerFont(TTFont(v[0], v[1]))

    return custom_fonts


def generate_pdf(data):
    """Generate a PDF document from namespace data.

    Content comes from generate_json(); this function only styles it.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, KeepTogether, Table, TableStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib import colors, fonts
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping
    from io import BytesIO
    from xml.sax.saxutils import escape
    import os

    custom_fonts = register_pdf_fonts()

    # Register font family mappings for bold/italic support in HTML tags
    # not works
    #addMapping('myfamily', 0, 0, 'Tinos-Regular')     # normal
    #addMapping('myfamily', 1, 0, 'Tinos-Bold')        # bold
    #addMapping('myfamily', 0, 1, 'Tinos-Italic')      # italic
    #addMapping('myfamily', 1, 1, 'Tinos-BoldItalic')  # bold-italic

    # Debug code - commented out
    # for (fam, b, i), fontName in fonts.font_index.items():
    #     if fam.lower() == family_to_check.lower():
    #         print(f"  - (Bold={b}, Italic={i}) -> '{fontName}'")
    #print(pdfmetrics.getRegisteredFontNames())

    buffer = BytesIO()

    # We'll use a custom document template to support two-column layout
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate

    doc = BaseDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )

    # Define frames for single and two-column layouts
    page_width, page_height = letter
    frame_width = page_width - 144  # Total width minus margins
    frame_height = page_height - 90  # Total height minus margins

    # Single column frame
    single_frame = Frame(
        72, 18, frame_width, frame_height,
        id='single_col',
        showBoundary=0
    )

    # Two column frames
    col_width = (frame_width - 20) / 2  # 20 points gap between columns
    left_frame = Frame(
        72, 18, col_width, frame_height,
        id='col_left',
        showBoundary=0
    )
    right_frame = Frame(
        72 + col_width + 20, 18, col_width, frame_height,
        id='col_right',
        showBoundary=0
    )

    # Add page templates
    doc.addPageTemplates([
        PageTemplate(id='SingleCol', frames=[single_frame]),
        PageTemplate(id='TwoCol', frames=[left_frame, right_frame]),
    ])

    # Container for the 'flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=custom_fonts['regular_tc'][0],
        fontSize=12,  # h2 size
        alignment=TA_CENTER,
        spaceAfter=8,
        spaceBefore=0,
    )

    author_style = ParagraphStyle(
        'CustomAuthor',
        parent=styles['Heading2'],
        fontName=custom_fonts['regular_tc'][0],
        fontSize=11,  # h3 size
        alignment=TA_CENTER,
        spaceAfter=8,
        spaceBefore=0,
    )

    section_heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName=custom_fonts['regular_tc'][0],
        fontSize=12,  # h3 size
        spaceAfter=4,
        spaceBefore=8,
    )
    category_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName=custom_fonts['bold_tc'][0],
        fontSize=12,  # h3 size
        spaceAfter=4,
        spaceBefore=12,
        alignment=TA_CENTER,
    )
    literature_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName=custom_fonts['regular_tc'][0],
        fontSize=10,  # h3 size
        spaceAfter=4,
        spaceBefore=12,
        alignment=TA_CENTER,
    )
    scientific_name_style = ParagraphStyle(
        'ScientificName',
        parent=styles['Normal'],
        fontName=custom_fonts['regular'][0],  # Tinos-Regular (base font)
        fontSize=10,
        spaceAfter=2,
        spaceBefore=0,
        leading=14,
    )

    body_style = ParagraphStyle(
        'BodyJustify',
        parent=styles['Normal'],
        fontName=custom_fonts['regular_tc'][0], #'Helvetica',  # Use Helvetica for HTML tag support
        alignment=TA_JUSTIFY,
        fontSize=10,
        spaceAfter=2,
        spaceBefore=4,
        leading=14,
        firstLineIndent=14,
    )

    normal_style = ParagraphStyle(
        'NormalText',
        parent=styles['Normal'],
        fontName=custom_fonts['regular_tc'][0],
        fontSize=10,
        spaceAfter=2,
        spaceBefore=0,
        leading=14,
    )

    # Add header (single column)
    from reportlab.platypus import NextPageTemplate



    # Build the presentation-free content context, then style it below
    context = generate_json(data)
    publications = context['publications']

    elements.append(NextPageTemplate('SingleCol'))
    elements.append(Paragraph(context['generator'], title_style))
    elements.append(Paragraph(f"Generated: {context['generatedAt']}", normal_style))
    elements.append(PageBreak())

    # Process each namespace
    for idx, d in enumerate(publications):
        # Add title and author (single column)
        elements.append(NextPageTemplate('SingleCol'))
        elements.append(Paragraph(d['title'], title_style))

        # Add item_category and literature section (single column)
        elements.append(Spacer(1, 0.2*inch))
        if cat := d.get('category'):
            elements.append(Paragraph(sanitize_html(cat['heading']), category_style))

            elements.append(Paragraph(d['author'], author_style))

            if desc := cat.get('description'):
                elements.append(Paragraph(sanitize_html(desc), body_style))

        elements.append(Paragraph('LITERATURE', literature_style))
        for content in d.get('literatures', []):
            elements.append(Paragraph(sanitize_html(content), normal_style))

        # Add identification keys if present (after literature)
        if keys := d.get('keys'):
            # Key entries style
            key_entry_style = ParagraphStyle(
                'KeyEntry',
                parent=styles['Normal'],
                fontName=custom_fonts['regular_tc'][0],
                fontSize=10,
                spaceAfter=2,
                spaceBefore=2,
                leading=14,
                leftIndent=0,
            )

            for key in keys:
                # Key title
                key_title = key.get('title', '')
                elements.append(Paragraph(f'檢索表: {key_title}', section_heading_style))
                elements.append(Spacer(1, 0.1*inch))

                # Render entries
                for entry in key.get('entries', []):
                    indent = '&nbsp;&nbsp;&nbsp;' * entry.get('indentLevel', 0)
                    number = entry.get('number', '')
                    description = sanitize_html(entry.get('description', ''))

                    # Item results are scientific names, so they are italicized
                    result_text = ''
                    if result := entry.get('result'):
                        if entry.get('resultType') == 'item':
                            result_text = f' ... {italicize_name(result)}'
                        else:
                            result_text = f' ... {result}'

                    entry_text = f'{indent}{number}. {description}{result_text}'
                    entry_text_with_fonts = convert_html_to_custom_fonts(entry_text, base_font='Serif')
                    elements.append(Paragraph(entry_text_with_fonts, key_entry_style))

                elements.append(Spacer(1, 0.2*inch))

        # Switch to two-column layout for species list
        elements.append(NextPageTemplate('TwoCol'))
        elements.append(PageBreak())

        # Process items (species) - will flow into two columns
        for v in d['items']:
            item_elements = []
            sci_name = f"<b>{v['number']}. {sanitize_html(v['scientificName'])}</b> {v['nameSuffix']}"

            # Convert HTML tags to font tags for custom font support
            sci_name_with_fonts = convert_html_to_custom_fonts(sci_name, base_font='Serif')
            item_elements.append(Paragraph(sci_name_with_fonts, scientific_name_style))

            # Common names
            if common := v.get('commonName'):
                normal_right_style = normal_style.clone('NormalRightText', alignment=TA_RIGHT, spaceBefore=12, spaceAfter=6)
                item_elements.append(Paragraph(sanitize_html(common), normal_right_style))

            # Synonyms
            for syn in v.get('synonyms', []):
                syn_with_fonts = convert_html_to_custom_fonts(sanitize_html(syn), base_font='Serif')
                item_elements.append(Paragraph(syn_with_fonts, scientific_name_style))

            # Description
            if desc := v.get('description'):
                item_elements.append(Paragraph(sanitize_html(desc), body_style))

            # Distribution
            if dist := v.get('distribution'):
                item_elements.append(Paragraph(sanitize_html(dist), body_style))

            # Specimens
            if specimens := v.get('specimens'):
                item_elements.append(Spacer(1, 0.1*inch))
                s = ''
                for group in specimens:
                    sp_str = '; '.join([sanitize_html(x) for x in group['records']])
                    s += f"{group['region']}: {sp_str}. "
                item_elements.append(Paragraph(s, normal_style))

            # Note
            if note := v.get('note'):
                item_elements.append(Paragraph(sanitize_html(note), body_style))

            # Add all item elements directly without KeepTogether to allow natural flow
            elements.extend(item_elements)
            # Small space between items (matching DOCX spacing)
            elements.append(Spacer(1, 6))

        # Add page break between namespaces (except for the last one)
        if idx < len(publications) - 1:
            elements.append(PageBreak())

    # Build PDF
    doc.build(elements)

    # Reset buffer position to beginning for reading
    buffer.seek(0)

    return buffer


# Metrics for the journal layout. generate_pdf2() renders JOURNAL_LAYOUT;
# generate_pdf3() renders JOURNAL_LAYOUT_COMPACT -- the same layout with
# smaller type, tighter leading and narrower margins, so more fits on a page.
# Type entries are (fontSize, leading) pairs; 'space' scales every
# spaceBefore / spaceAfter in the styles.
JOURNAL_LAYOUT = {
    # page
    'margin_mm': 18,
    'head_reserve': 30,       # running head strip above the content area
    'foot_reserve': 24,       # running footer strip below the content area
    'foot_rule_offset': 15,   # footer rule, above the bottom margin
    'col_gap': 20,
    'min_body_height': 140,   # keep at least this much of page 1 for the columns
    'head_slack': 6,          # so the front block never spills into the columns
    'head_gap': 12,           # between the front block and the columns
    'running_title_max': 70,
    # front block: what is shown, and how it is aligned
    'show_title': True,       # FM-TITLE; when off it survives in the running head
    'taxon_align': 'left',    # FM-TAXON:  left | center
    'taxon_weight': 'bold',   # FM-TAXON:  bold | regular
    'taxon_authors_font': None,  # FM-TAXON.authors: font role, or None for
                                 # the same weight as the rest of the heading
    'author_align': 'left',   # FM-AUTHOR: left | center
    # type: (fontSize, leading)
    'head_font': 8.5,
    'foot_font': 8,
    'title': (21, 27),
    'taxon_heading': (14, 20),
    'author': (11.2, 17),
    'description': (10.1, 17),
    'section': (12.4, 16),
    'taxon': (11.3, 15),
    'common': (9.8, 14),
    'synonym': (9.8, 14),
    'body': (10.9, 17.6),
    'key': (10.2, 15),
    'reference': (9.4, 15),
    # spacing and indents
    'space': 1.0,
    'indent': 10.5,           # hanging indent, synonyms and references
    'key_indent': 10,         # per indentLevel
    'key_number_w': 15,
    'key_pad': 5,             # row padding in the key table
}

JOURNAL_LAYOUT_COMPACT = {
    **JOURNAL_LAYOUT,
    'margin_mm': 15,
    'head_reserve': 26,
    'foot_reserve': 20,
    'foot_rule_offset': 12,
    'col_gap': 16,
    'min_body_height': 120,
    'head_slack': 5,
    'head_gap': 10,
    'running_title_max': 82,
    'show_title': False,
    'taxon_align': 'center',
    'taxon_weight': 'regular',
    'taxon_authors_font': 'medium_tc',
    'author_align': 'center',
    'head_font': 7.5,
    'foot_font': 7,
    'title': (16.5, 21),
    'taxon_heading': (13.5, 18),
    'author': (9.5, 13.5),    # below the heading, and stays bold
    'description': (8.8, 13.6),
    'section': (10.4, 13.5),
    'taxon': (9.6, 12.8),
    'common': (8.4, 11.6),
    'synonym': (8.4, 11.6),
    'body': (9.2, 13.8),
    'key': (8.8, 12.4),
    'reference': (8.2, 12),
    'space': 0.8,
    'indent': 9,
    'key_indent': 8,
    'key_number_w': 13,
    'key_pad': 3.5,
}


def generate_pdf2(data):
    """Generate a journal-article style PDF ("Biota Journal" layout).

    Same content as generate_pdf() -- both render generate_json() -- but laid
    out as a journal article: A4/18mm page with a running head and footer, a
    full-width front block (section 1: centered title, taxon heading, author and
    description; section 2: LITERATURE list), then a two-column body with
    the species treatments and the identification keys. The design's specimen
    data table and figure slots are intentionally not rendered.
    """
    return build_journal_pdf(data, JOURNAL_LAYOUT)


def generate_pdf3(data):
    """Generate the compact variant of the journal layout.

    Same structure and content as generate_pdf2(), rendered with
    JOURNAL_LAYOUT_COMPACT: smaller type throughout, tighter leading and
    spacing, 15mm margins and a narrower column gap, so a treatment takes
    noticeably less space on the page.
    """
    return build_journal_pdf(data, JOURNAL_LAYOUT_COMPACT)


def build_journal_pdf(data, layout):
    """Render the journal layout with the given metrics.

    `layout` is a JOURNAL_LAYOUT-shaped dict: it carries every page dimension,
    font size and spacing this renderer uses, so a variant only has to override
    the numbers it changes. See generate_pdf2() / generate_pdf3().
    """
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageTemplate, Paragraph, Table,
        TableStyle, PageBreak, FrameBreak, NextPageTemplate,
    )
    from reportlab.platypus.flowables import HRFlowable

    custom_fonts = register_pdf_fonts()
    F_TC = custom_fonts['regular_tc'][0]
    F_TC_TEXT = custom_fonts['text_tc'][0]
    F_TC_BOLD = custom_fonts['bold_tc'][0]
    F_AUTHORS = (custom_fonts[layout['taxon_authors_font']][0]
                 if layout['taxon_authors_font'] else None)

    # palette taken from the journal layout design
    DARK = HexColor('#1A1A1A')
    TEXT = HexColor('#333333')
    MUTED = HexColor('#767676')
    RULE = HexColor('#E6E6E6')
    RULE_STRONG = HexColor('#949494')

    page_w, page_h = A4
    MARGIN = layout['margin_mm'] * mm
    HEAD_RESERVE = layout['head_reserve']
    FOOT_RESERVE = layout['foot_reserve']
    COL_GAP = layout['col_gap']
    MIN_BODY_HEIGHT = layout['min_body_height']
    SP = layout['space']
    INDENT = layout['indent']
    ALIGN = {'left': TA_LEFT, 'center': TA_CENTER, 'justify': TA_JUSTIFY}

    frame_w = page_w - 2 * MARGIN
    col_w = (frame_w - COL_GAP) / 2
    content_top = page_h - MARGIN - HEAD_RESERVE
    content_bottom = MARGIN + FOOT_RESERVE
    usable_h = content_top - content_bottom
    frame_pad = {
        'leftPadding': 0,
        'rightPadding': 0,
        'topPadding': 0,
        'bottomPadding': 0,
        'showBoundary': 0,
    }

    context = generate_json(data)
    publications = context['publications']

    def plain(text):
        """Strip inline markup, for text drawn directly on the canvas."""
        return re.sub(r'<[^>]+>', '', text or '').strip()

    def markup(text):
        """Prepare inline markup (<i>, <b>, <br>) for the custom fonts."""
        return convert_html_to_custom_fonts(sanitize_html(text))

    def bold(text):
        return f'<font name="{F_TC_BOLD}">{text}</font>'

    running_title = plain(publications[0]['title']) if publications else ''
    title_max = layout['running_title_max']
    if len(running_title) > title_max:
        running_title = f'{running_title[:title_max - 1]}…'

    def draw_furniture(canvas, doc_):
        """Running head and footer, drawn on every page."""
        canvas.saveState()

        head_font = layout['head_font']
        top = page_h - MARGIN
        canvas.setFont(F_TC_BOLD, head_font)
        canvas.setFillColor(DARK)
        canvas.drawString(MARGIN, top - (head_font + 0.5), context['generator'])
        canvas.setFont(F_TC, head_font)
        canvas.setFillColor(MUTED)
        canvas.drawRightString(page_w - MARGIN, top - (head_font + 0.5), running_title)
        canvas.setStrokeColor(RULE_STRONG)
        canvas.setLineWidth(0.6)
        rule_y = top - (head_font + 6.5)
        canvas.line(MARGIN, rule_y, page_w - MARGIN, rule_y)

        foot_font = layout['foot_font']
        base = MARGIN + layout['foot_rule_offset']
        canvas.setStrokeColor(RULE)
        canvas.setLineWidth(0.6)
        canvas.line(MARGIN, base, page_w - MARGIN, base)
        canvas.setFont(F_TC, foot_font)
        canvas.setFillColor(MUTED)
        canvas.drawString(MARGIN, base - (foot_font + 2), context['generatedAt'])
        canvas.drawRightString(page_w - MARGIN, base - (foot_font + 2), str(doc_.page))

        canvas.restoreState()

    # --- styles ---------------------------------------------------------
    title_style = ParagraphStyle(
        'JTitle', fontName=F_TC_BOLD, fontSize=layout['title'][0],
        leading=layout['title'][1], textColor=DARK,
        alignment=TA_CENTER, spaceBefore=2 * SP, spaceAfter=10 * SP)
    author_style = ParagraphStyle(
        'JAuthor', fontName=F_TC_BOLD, fontSize=layout['author'][0],
        leading=layout['author'][1], textColor=DARK,
        alignment=ALIGN[layout['author_align']], spaceAfter=2 * SP)
    description_style = ParagraphStyle(
        'JDescription', fontName=F_TC, fontSize=layout['description'][0],
        leading=layout['description'][1], textColor=TEXT,
        alignment=TA_JUSTIFY, spaceBefore=6 * SP, spaceAfter=7 * SP)
    # the unbold heading uses the 400-weight face, not the Light one, so the
    # italic scientific name and the upright authors carry the same weight
    taxon_bold = layout['taxon_weight'] == 'bold'
    taxon_heading_style = ParagraphStyle(
        'JTaxonHeading', fontName=F_TC_BOLD if taxon_bold else F_TC_TEXT,
        fontSize=layout['taxon_heading'][0],
        leading=layout['taxon_heading'][1], textColor=DARK,
        alignment=ALIGN[layout['taxon_align']],
        spaceBefore=2 * SP, spaceAfter=6 * SP)
    section_style = ParagraphStyle(
        'JSection', fontName=F_TC_BOLD, fontSize=layout['section'][0],
        leading=layout['section'][1], textColor=DARK,
        spaceBefore=12 * SP, spaceAfter=6 * SP, keepWithNext=1)
    literature_heading_style = ParagraphStyle(
        'JLiteratureHeading', parent=section_style, alignment=TA_CENTER)
    taxon_style = ParagraphStyle(
        'JTaxon', fontName=F_TC, fontSize=layout['taxon'][0],
        leading=layout['taxon'][1], textColor=DARK,
        spaceBefore=9 * SP, spaceAfter=2 * SP, keepWithNext=1)
    common_style = ParagraphStyle(
        'JCommon', fontName=F_TC, fontSize=layout['common'][0],
        leading=layout['common'][1], textColor=MUTED,
        spaceAfter=4 * SP)
    synonym_style = ParagraphStyle(
        'JSynonym', fontName=F_TC, fontSize=layout['synonym'][0],
        leading=layout['synonym'][1], textColor=TEXT,
        leftIndent=INDENT, firstLineIndent=-INDENT, spaceAfter=1 * SP)
    body_style = ParagraphStyle(
        'JBody', fontName=F_TC, fontSize=layout['body'][0],
        leading=layout['body'][1], textColor=TEXT,
        alignment=TA_JUSTIFY, spaceBefore=4 * SP, spaceAfter=6 * SP)
    key_number_style = ParagraphStyle(
        'JKeyNumber', fontName=F_TC_BOLD, fontSize=layout['key'][0],
        leading=layout['key'][1], textColor=DARK)
    key_text_style = ParagraphStyle(
        'JKeyText', fontName=F_TC, fontSize=layout['key'][0],
        leading=layout['key'][1], textColor=TEXT)
    reference_style = ParagraphStyle(
        'JReference', fontName=F_TC, fontSize=layout['reference'][0],
        leading=layout['reference'][1], textColor=TEXT,
        leftIndent=INDENT, firstLineIndent=-INDENT, spaceAfter=5 * SP)

    # --- title block (full page width) ----------------------------------
    def build_header(pub):
        flow = []

        # section 1: title, taxon heading, author, description
        if layout['show_title']:
            flow.append(Paragraph(markup(pub['title']), title_style))

        cat = pub.get('category')
        if cat and (heading := cat.get('heading')):
            # italicize the leading scientific name, keep the rest upright,
            # and optionally set the name-authors in their own weight
            name = cat.get('scientificName', '')
            authors = cat.get('authors', '')
            common = cat.get('commonNames', '')
            if heading == ' '.join(x for x in (name, authors, common) if x):
                # the documented "{sci} {authors} {common}" join: rebuild it
                # part by part, so each part can be styled on its own
                parts = [italicize_name(name)] if name else []
                if authors:
                    parts.append(f'<font name="{F_AUTHORS}">{authors}</font>'
                                 if F_AUTHORS else authors)
                if common:
                    parts.append(common)
                heading = ' '.join(parts)
            elif name and heading.startswith(name):
                heading = f'{italicize_name(name)}{heading[len(name):]}'
            text = markup(heading)
            flow.append(Paragraph(
                bold(text) if taxon_bold else text, taxon_heading_style))

        if author := pub.get('author'):
            flow.append(Paragraph(markup(author), author_style))

        if cat and (desc := cat.get('description')):
            flow.append(Paragraph(markup(desc), description_style))

        # section 2: literature list, full width
        if literatures := pub.get('literatures'):
            flow.append(Paragraph('LITERATURE', literature_heading_style))
            for content in literatures:
                flow.append(Paragraph(markup(content), reference_style))
            flow.append(HRFlowable(
                width='100%', thickness=0.6, color=RULE,
                spaceBefore=6 * SP, spaceAfter=0))

        return flow

    # --- two-column body ------------------------------------------------
    def build_body(pub):
        flow = []

        # the taxon heading and its description live in the title block
        for v in pub['items']:
            name = f"{bold(str(v['number']) + '.')} {italicize_name(v['scientificName'])}"
            if suffix := v.get('nameSuffix'):
                name = f'{name} {suffix}'
            flow.append(Paragraph(markup(name), taxon_style))

            if common := v.get('commonName'):
                flow.append(Paragraph(markup(common), common_style))

            for syn in v.get('synonyms', []):
                flow.append(Paragraph(markup(syn), synonym_style))

            if desc := v.get('description'):
                flow.append(Paragraph(markup(desc), body_style))

            if dist := v.get('distribution'):
                flow.append(Paragraph(
                    f'{bold("Distribution.")} {markup(dist)}', body_style))

            if specimens := v.get('specimens'):
                groups = []
                for group in specimens:
                    records = '; '.join(sanitize_html(x) for x in group['records'])
                    groups.append(f"{group['region']}: {records}.")
                flow.append(Paragraph(
                    f'{bold("Material examined.")} {markup(" ".join(groups))}',
                    body_style))

            if note := v.get('note'):
                flow.append(Paragraph(f'{bold("Note.")} {markup(note)}', body_style))

        for key in pub.get('keys', []):
            title = key.get('title') or ''
            flow.append(Paragraph(
                f'檢索表 Key{": " + title if title else ""}', section_style))

            rows = []
            for entry in key.get('entries', []):
                result = ''
                if r := entry.get('result'):
                    result = (f' … {italicize_name(r)}'
                              if entry.get('resultType') == 'item' else f' … {r}')
                indent = entry.get('indentLevel', 0) * layout['key_indent']
                text_style = key_text_style
                if indent:
                    text_style = key_text_style.clone(
                        f'JKeyText{indent}', leftIndent=indent)
                rows.append([
                    Paragraph(str(entry.get('number', '')), key_number_style),
                    Paragraph(
                        markup(f"{sanitize_html(entry.get('description', ''))}{result}"),
                        text_style),
                ])

            if rows:
                number_w = layout['key_number_w']
                table = Table(rows, colWidths=[number_w, col_w - number_w - 0.5])
                table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), layout['key_pad']),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), layout['key_pad']),
                    ('LINEABOVE', (0, 0), (-1, 0), 0.6, RULE_STRONG),
                    ('LINEBELOW', (0, 0), (-1, -1), 0.5, RULE),
                ]))
                flow.append(table)

        return flow

    def measure(flowables, width):
        """Height the title block needs, so the columns can start below it."""
        total = 0.0
        pending = 0.0
        for f in flowables:
            try:
                _, h = f.wrap(width, usable_h)
                space_before, space_after = f.getSpaceBefore(), f.getSpaceAfter()
            except Exception:
                return None
            total += max(space_before, pending) + h
            pending = space_after
        return total + pending

    # --- page templates: one title page per publication, plus columns ---
    blocks = []
    templates = []
    for idx, pub in enumerate(publications):
        header = build_header(pub)
        body = build_body(pub)

        header_h = measure(header, frame_w)
        # if the title block cannot leave room for a usable column pair,
        # give it a page of its own and start the body on the next page
        own_page = header_h is None or header_h > usable_h - MIN_BODY_HEIGHT
        if own_page:
            frames = [Frame(MARGIN, content_bottom, frame_w, usable_h,
                            id=f'head{idx}', **frame_pad)]
        else:
            header_h += layout['head_slack']
            body_h = usable_h - header_h - layout['head_gap']
            frames = [
                Frame(MARGIN, content_top - header_h, frame_w, header_h,
                      id=f'head{idx}', **frame_pad),
                Frame(MARGIN, content_bottom, col_w, body_h,
                      id=f'first_left{idx}', **frame_pad),
                Frame(MARGIN + col_w + COL_GAP, content_bottom, col_w, body_h,
                      id=f'first_right{idx}', **frame_pad),
            ]

        templates.append(PageTemplate(
            id=f'Head{idx}', frames=frames, onPage=draw_furniture))
        blocks.append((header, body, own_page))

    templates.append(PageTemplate(id='TwoCol', frames=[
        Frame(MARGIN, content_bottom, col_w, usable_h, id='col_left', **frame_pad),
        Frame(MARGIN + col_w + COL_GAP, content_bottom, col_w, usable_h,
              id='col_right', **frame_pad),
    ], onPage=draw_furniture))

    buffer = BytesIO()
    doc = BaseDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=MARGIN + HEAD_RESERVE,
        bottomMargin=MARGIN + FOOT_RESERVE,
        title=running_title,
        author=publications[0]['author'] if publications else '',
    )
    doc.addPageTemplates(templates)

    elements = []
    for idx, (header, body, own_page) in enumerate(blocks):
        if idx:
            elements.append(NextPageTemplate(f'Head{idx}'))
            elements.append(PageBreak())

        elements.extend(header)
        if own_page:
            elements.append(NextPageTemplate('TwoCol'))
            elements.append(PageBreak())
        else:
            elements.append(FrameBreak())
            elements.append(NextPageTemplate('TwoCol'))
        elements.extend(body)

    doc.build(elements)
    buffer.seek(0)

    return buffer


def get_namespace_data(namespace_id):
    conn = get_db_connection()
    data = {
        'title': '',
        'author': '',
        'literatures': [],
        'items': [],
        'id': namespace_id,
        'reference_name': ''
    }

    #mysql_cursor.execute(literature_sql)
    mysql_cursor = conn.cursor()
    #with conn.cursor() as cursor:
    literature_sql = f'SELECT a.author, a.reference_id, a.content FROM import_checklist_logs c LEFT JOIN api_citations a ON FIND_IN_SET(a.reference_id, c.included_references) > 0 WHERE c.namespace_id = {namespace_id}'
    mysql_cursor.execute(literature_sql)
    rows = mysql_cursor.fetchall()
    for r in rows:
        data['literatures'].append({'author': r['author'], 'id': r['reference_id'], 'content': r['content']})

    mysql_cursor.execute(f'SELECT n.title, u.name FROM my_namespaces n LEFT JOIN users u ON u.id = n.user_id WHERE n.id={namespace_id}')
    result = mysql_cursor.fetchone()
    data['title'] = result['title']
    data['author'] = result['name']

    mysql_cursor.execute(f"SELECT t.name, t._authorship, t.id, u.per_usages, u.type_specimens, u.properties, r.title, u.id, t.properties, u.name_remark, u.group, u.updated_at FROM my_namespace_usages u LEFT JOIN taxon_names t ON u.taxon_name_id = t.id LEFT JOIN `references` r ON r.id = t.reference_id WHERE namespace_id={namespace_id} AND u.status='accepted' ORDER by `order`")

    rows = mysql_cursor.fetchall()
    for row in rows:
        #print(row)
        common_names = []
        synonyms = []
        note = ''
        description = ''
        distribution = ''
        add_fields = {}
        specimens = []
        taicol_name_id = row['id']

        per_usages = []
        type_specimens = {}
        properties = {}

        if x := row['per_usages']:
            #source_data['usages'] = yaml.dump(json.loads(x), default_flow_style=False, sort_keys=False, allow_unicode=True)
            per_usages = json.loads(row['per_usages'])
        if x := row['properties']:
            #description:"特徵描述",diagnosis:"鑑定特徵",distribution:"物種分布",etymology:"語源",habitat:"棲地",substrata:"基質",measurements:"測量",coloration:"顏色",otherExaminedMaterial:"其他引證標本"
            properties = json.loads(x)
            #source_data['props'] = yaml.dump(props, default_flow_style=False, sort_keys=False, allow_unicode=True)

            if names := properties.get('common_names'):
                for n in names:
                    if x := n['name']:
                        common_names.append(x.replace('\u0000', '').strip())

        if x := row['type_specimens']:
            type_specimens = json.loads(x)
            #source_data['type'] = yaml.dump(sp, default_flow_style=False, sort_keys=False, allow_unicode=True)

        #if x := row[6]:
        #    if x not in data['literatures']:
        #        data['literatures'].append(x)

        if x:= row['group']:
            # mysql_cursor.execute(f"SELECT ru.id, ru.status, t.name, t._authorship, r.title FROM reference_usages ru LEFT JOIN `references` r ON r.id = ru.reference_id LEFT JOIN taxon_names t ON t.id = ru.taxon_name_id WHERE ru.accepted_taxon_name_id={x} and ru.status != 'accepted'")
            # results = mysql_cursor.fetchall()
            # sci_name = ''
            # ref_title = ''
            # for i in results:
            #     sci_name = i['name']
            #     if author := i['_authorship']:
            #         sci_name = f'{sci_name} {author}'
            #     if ref := i['title']:
            #         ref_title = ref
            #     synonyms.append([sci_name, ref_title])

            mysql_cursor.execute(f"SELECT t.name, t.formatted_authors, t.id, u.properties FROM my_namespace_usages u LEFT JOIN taxon_names t ON u.taxon_name_id = t.id WHERE u.namespace_id={namespace_id} AND u.status != 'accepted' AND u.group = {x}")

            res_synonyms = mysql_cursor.fetchall()
            for x in res_synonyms:
                y = {}
                for k, v in x.items():
                    if k == 'properties':
                        y[k] = json.loads(v)
                    else:
                        y[k] = v
                synonyms.append(y)

        # item_title atomic struct
        item_title = {
            'scientific_name': {
                'canonical': row['name'],
                'author': '',
                'full': ''
            },
            'ref': '',
            'name_in_ref': '',
            'type': {
                'use': '',
                'gathering': {
                    'country': '',
                    'locality': '',
                    'locality2': '',
                    'year': '',
                    'month': '',
                    'day': '',
                    'field_number': '',
                },
                'specimens': [],
            },
            'description': description,
            'distributions': distribution,
            #'display_name': ['', ['', '', ''],'',''], # ['canonical_name', ["author", "ref", "name-in-ref"], "type status:"]
        }

        if x := row['name']:
            item_title['scientific_name']['author'] = row['_authorship']
            item_title['scientific_name']['full'] = f"{item_title['scientific_name']['canonical']} {x}"
            #item_title['display_name'][1][0] = x

        if x := row['properties']:
            taxon_props = json.loads(x)
            if y := taxon_props.get('reference_name'):
                item_title['ref'] = y
                #item_title['display_name'][1][1] = y

        for u in per_usages:
            if x := u.get('name_in_reference'):
                item_title['name_in_ref'] = x
                #item_title['display_name'][1][2] = x

        for s in type_specimens:
            if x := s.get('use'):
                item_title['type']['use'] = x

            # gathering
            if country := s.get('country'):
                if d := country.get('display'):
                    if c := d.get('en-us'):
                        item_title['type']['gathering']['country'] = c

            if x := s.get('locality'):
                item_title['type']['gathering']['locality'] = x
            if x := s.get('locality_verbatim'):
                item_title['type']['gathering']['locality2'] = x

            if x := s.get('collection_year'):
                item_title['type']['gathering']['year'] = x
            if x := s.get('collection_month'):
                item_title['type']['gathering']['month'] = x
            if x := s.get('collection_day'):
                item_title['type']['gathering']['day'] = x

            if x := s.get('specimens'):
                item_title['type']['specimens'] = x


        # format display string
        #tmp = f"{item_title['scientific_name']}"
        tmp = ''
        if x := item_title['scientific_name']['author']:
            tmp = x
        if x := item_title['ref']:
            tmp = f"{tmp}, {x}"
        if x := item_title['name_in_ref']:
            if tmp[-1] == '.':
                tmp = tmp[:-1] # Berberis aristatoserrulata, ref 最後有一個. => 組起來要拿掉
            tmp = f"{tmp}, {x}."

        voucher = ''
        if x:= item_title['type']['use']:
            voucher = x.capitalize()
        if x := item_title['type']['gathering']['country']:
            voucher = f"{voucher}: {x.upper()}"
        loc = ''
        if x := item_title['type']['gathering']['locality']:
            loc = x
            if x2 := item_title['type']['gathering']['locality2']:
                loc = f'{loc}({x2})'
        if loc:
            voucher = f'{voucher}, {loc}'

        ymd = []
        if x := item_title['type']['gathering']['day']:
            ymd.append(x)
        if x := item_title['type']['gathering']['month']:
            ymd.append(x)
        if x := item_title['type']['gathering']['year']:
            ymd.append(x)
        if ymd:
            voucher = f"{voucher}, {' '.join(ymd)}"

        if x := item_title['type']['gathering']['field_number']:
            vouche = f"{voucher}, {x}"
        #else:
        #    voucher = f"{voucher}, sine coll" # s.n.

        sp = []
        for index, x in enumerate(item_title['type']['specimens']):
            s = x['herbarium']
            prefix = ''
            if index > 0:
                prefix = 'isotype: '

            if an := x.get('accession_number'):
                sp.append(f"{prefix}{s} [{an}]")
            else:
                sp.append(f"{prefix}{s}")

        if len(sp):
            voucher = f"{voucher}. ({'; '.join(sp)})."

        if voucher:
            tmp = f"{tmp} {voucher}"

        # parse name_remark
        full_name = ''
        if name_remark := row['name_remark']:
            soup = BeautifulSoup(name_remark, 'lxml')
            if soup.i:
                name_remark = soup.i.string
                full_name = soup.get_text()
                full_name = full_name.replace(name_remark, '').strip()

        #print('---')
        #print(full_name)
        #print(tmp)
        #print(item_title)
        item_title['name_remark'] = full_name
        item_title['voucher'] = voucher
        item_title['display_name'] = tmp

        data['items'].append({
            'item_title': item_title,
            'status': row['id'],
            'commonNames': common_names,
            'synonyms': synonyms,
            'type_specimens': type_specimens,
            'properties': properties,
            'taicol_taxon_name_id': taicol_name_id,
            'taicol_usage_id': row['id'],
            'updated': row['updated_at'].strftime('%Y-%m-%d %H:%M:%S'),
        })

    return data

class TBIASpecimens(object):

    records = []
    url = ''
    def __init__(self):

        pass

    def fetch_taxon(self, taxon_key):
        self.url = f'https://tbiadata.tw/api/v1/occurrence?isCollection=true&taxonID={taxon_key}&limit=100'
        current_app.logger.debug(f'TBIASpecimens) fetch_taxon: {taxon_key}')
        self.records = []
        res = self.fetch_api(self.url)
        return {
            'is_success': res['is_success'],
            'records': self.records,
        }


    def fetch_api(self, url):
        current_app.logger.debug(f'TBIASpecimens) fetch_api: {url}')
        resp = requests.get(url)
        is_success = False
        try:
            data = resp.json()
            if data['status']['code']!= 200:
                is_success = False
            else:
                self.records += data['data']
                is_success = True

            if next_url := data['links'].get('next'):
                self.fetch_api(next_url)
        except:
            current_app.logger.error('TBIASpecimens) fetch_api error')


        return {
            'is_success': is_success,
            'records': self.records,
        }


    def conv_data(self, data):
        DWC_TERMS = {
            'countryCode': 'cc',
            'stateProvince': 'adm1',
            'county': 'adm2',
            'municipality': 'adm3',
            'locationID': '',
            'higherGeographyID': '',
            'higherGeography': '',
            'continent': '',
            'waterBody': '',
            'islandGroup': '',
            'island': '',
            'country': '',
            'locality': '',
            'verbatimLocality': '',
            'minimumElevationInMeters': '',
            'maximumElevationInMeters': '',
            'verbatimElevation': '',
            'verticalDatum': '',
            'minimumDepthInMeters': '',
            'maximumDepthInMeters': '',
            'verbatimDepth': '',
            'minimumDistanceAboveSurfaceInMeters': '',
            'maximumDistanceAboveSurfaceInMeters': '',
            'locationAccordingTo': '',
            'locationRemarks': '',
            'decimalLatitude': '',
            'decimalLongitude': '',
            'geodeticDatum': '',
            'coordinateUncertaintyInMeters': '',
            'coordinatePrecision': '',
            'pointRadiusSpatialFit': '',
            'verbatimCoordinates': '',
            'verbatimLatitude': '',
            'verbatimLongitude': '',
            'verbatimCoordinateSystem': '',
            'verbatimSRS': '',
            'footprintWKT': '',
            'footprintSRS': '',
            'footprintSpatialFit': '',
            'georeferencedBy': '',
            'georeferencedDate': '',
            'georeferenceProtocol': '',
            'georeferenceSources': '',
            'georeferenceRemarks': '',
        }

        TBIA_TERMS = {
            'county': 'adm1',
            'municipality': 'adm2',
        }

        taiwan_counties_english = {
            '宜蘭縣': 'Yilan',
            '桃園市': 'Taoyuan',
            '新竹縣': 'Hsinchu',
            '苗栗縣': 'Miaoli',
            '彰化縣': 'Changhua',
            '南投縣': 'Nantou',
            '雲林縣': 'Yunlin',
            '嘉義縣': 'Chiayi',
            '屏東縣': 'Pingtung',
            '臺東縣': 'Taitung',
            '花蓮縣': 'Hualien',
            '澎湖縣': 'Penghu',
            '基隆市': 'Keelung',
            '新竹市': 'Hsinchu',
            '嘉義市': 'Chiayi',
            '臺北市': 'Taipei',
            '新北市': 'New Taipei',
            '臺中市': 'Taichung',
            '臺南市': 'Tainan',
            '高雄市': 'Kaohsiung'
        }
        records = []
        for i,v in enumerate(data):
            specimen_display = {
                'county': '',
                'locality': '',
                'collector': '',
                'record_number': '',
            }
            named_areas = {}
            for term, field in DWC_TERMS.items():
                if x := v.get(term):
                    key = term if field == '' else field
                    named_areas[key] = x
                    if key == 'adm2':
                        if county_en := taiwan_counties_english.get(x):
                            specimen_display['county'] = county_en.upper()
                    if key == 'locality':
                        specimen_display['locality'] = x
                        #'ILAN: Nanhutashan, Lu 24973.
                        # NANTOU: Mt.Kiraishiu, Wilson 10074 (Type of B. nantoensis, A!);'
            media = []
            if x := v.get('associatedMedia'):
                media.append(x)

            locality_list = []
            if x:= v.get('county'):
                locality_list.append(x)
            if x:= v.get('locality'):
                locality_list.append(x)

            date = ''
            if x := v.get('eventDate'):
                y = x[0:4]
                m = x[5:7]
                d = x[8:10]
                date = f'{y}.{m}.{d}'

            recorded_by = ''
            if x := v.get('recordedBy', ''):
                recorded_by = x
                specimen_display['collector'] = recorded_by
                record_number = ''
            if x := v.get('recordNumber', ''):
                record_number = x
                specimen_display['record_number'] = x

            records.append({
                'recid': i,
                'url': v.get('references', ''),
                'institutionCode': '',
                'recordedBy': recorded_by,
                'basisOfRecord': v.get('basisOfRecord', ''),
                'recordNumber': record_number,
                'catalogNumber': v.get('catalogNumber', ''),
                'date': date,
                'remarks': '',
                'locality': '|'.join(locality_list),
                'datasetTitle': v['datasetName'],
                'media': media,
                'named_areas': named_areas,
                'specimen_display': specimen_display,
                '_raw': v,
            })

        return records


def send_email(to, subject, body):
    ses_client = boto3.client(
        'ses',
        region_name=current_app.config['AWS_SES_REGION_NAME'],
        aws_access_key_id=current_app.config['AWS_SES_ACCESS_KEY'],
        aws_secret_access_key=current_app.config['AWS_SES_SECRET_ACCESS_KEY']
    )

    try:
        response = ses_client.send_email(
            Source=current_app.config['AWS_SES_SOURCE'],
            Destination={
                'ToAddresses': [
                    to,
                ],
            },
            Message={
                'Subject': {
                    'Data': subject,
                    'Charset': 'UTF-8'
                },
                'Body': {
                    'Text': {
                        'Data': body,
                        'Charset': 'UTF-8'
                    }
                }
            }
        )

        current_app.logger.info(f"Email sent successfully! Message ID: {response['MessageId']}")
        return response

    except ClientError as e:
        error_code = e.response['Error']['Code']
        error_message = e.response['Error']['Message']
        current_app.logger.error(f"Error sending email: {error_code} - {error_message}")
        return None


def put_publication_by_taicol_namespace(user, namespace_id):
    result = {
        'message': '',
        'status': '',
        'data': {},
    }

    if is_ok := check_namespace_available(user.email, namespace_id):
        publication_id = None
        message = ''
        if pub := Publication.query.join(Collection).filter(
            Collection.user_id==user.id,
            Collection.source_id==namespace_id,
            Collection.source_name=='taicol:namespace').scalar(): # exist
            message = f'publication [{pub.id}] already exist'
            publication_id = pub.id
            result['status'] = 'exist'
        else:
            result['status'] = 'new'
            pub = Publication(title='', author='')
            session.add(pub)
            session.commit()

            c = create_collection_by_taicol_namespace(user.id, namespace_id)
            pub.title = c.name
            pub.author = c.source_data.get('author')
            c.publication_id = pub.id

            for i, v in enumerate(c.source_data['literatures']):
                pl = PublicationLiterature(publication_id=pub.id, source_id=v['reference_id'], name=v['citation'], sort=i+1)
                session.add(pl)

            session.commit()
            publication_id = pl.id

        current_app.logger.info(message)
        result.update({
            'message': message,
            'data': {
                'publication_id': publication_id,
            }
        })
        return result

    return result


def create_collection_by_taicol_namespace(user_id, namespace_id):
    url = f"{current_app.config['TAICOL_API']}/biota?namespace_id={namespace_id}&token={current_app.config['TAICOL_TOKEN']}"
    resp = requests.get(url)
    data = resp.json()
    collection = Collection(name=data['title'], source_name=f'taicol:namespace', source_data=data, source_id=namespace_id, user_id=user_id)
    session.add(collection)
    session.commit()

    counter = 0
    for i in data['group']:
        counter += 1
        name = i['name'].replace('<i>', '').replace('</i>', '')
        item = Item(
            collection_id=collection.id,
            description=i['description'],
            distribution=i['distribution'],
            note=i['note'],
            user_id=user_id,
            scientific_name=name,
            source_data=i,
            common_names='|'.join(i['common_names']),
            sort=counter,
        )
        session.add(item)
        session.commit()

        for syn in i['synonyms']:
            item_syn = ItemSynonym(item_id=item.id, name=syn['usage_references_text'], ref=f"name_id:{i['name_id']}")
            session.add(item_syn)

        session.commit()

    return collection


def sync_collection_by_taicol_namespace(collection):
    """Update existing collection and items with latest data from TaiCOL"""
    namespace_id = collection.source_id
    url = f"{current_app.config['TAICOL_API']}/biota?namespace_id={namespace_id}&token={current_app.config['TAICOL_TOKEN']}"
    resp = requests.get(url)
    data = resp.json()

    # Update collection
    collection.name = data['title']
    collection.source_data = data

    remote_items = [i['name_id'] for i in data['group']]
    # Build existing items lookup by name_id
    existing_items = {}
    counter_del = 0
    counter_new = 0
    counter_edit = 0

    for item in collection.items:
        source_id = item.source_data.get('name_id')
        if item.source_data and source_id:
            existing_items[source_id] = item
            if source_id not in remote_items:
                counter_del += 1
                for syn in item.synonyms:
                    session.delete(syn)
                for spec in item.specimens:
                    session.delete(spec)
                for img in item.images:
                    session.delete(img)
                for dist in item.distributions:
                    session.delete(dist)
                session.delete(item)

    session.commit()
    counter = 0

    for i in data['group']:
        counter += 1
        name = i['name'].replace('<i>', '').replace('</i>', '')
        name_id = i.get('name_id')

        if name_id and name_id in existing_items:
            counter_edit += 1
            # Update existing item
            item = existing_items[name_id]
            item.scientific_name = name
            item.description = i['description']
            item.distribution = i['distribution']
            item.note = i['note']
            item.source_data = i
            item.common_names = '|'.join(i['common_names'])
            item.sort = counter

            # Update synonyms: delete old and create new
            for syn in item.synonyms:
                session.delete(syn)
        else:
            # Create new item
            item = Item(
                collection_id=collection.id,
                description=i['description'],
                distribution=i['distribution'],
                note=i['note'],
                user_id=collection.user_id,
                scientific_name=name,
                source_data=i,
                common_names='|'.join(i['common_names']),
                sort=counter
            )
            session.add(item)
            session.commit()
            counter_new += 1

        # Add synonyms
        for syn in i['synonyms']:
            item_syn = ItemSynonym(item_id=item.id, name=syn['usage_references_text'], ref=f"name_id:{i['name_id']}")
            session.add(item_syn)

    session.commit()
    return {
        'deleted': counter_del,
        'created': counter_new,
        'updated': counter_edit,
    }

def format_specimen_display(data):
    record_number = data.get('recordNumber', '--')
    recorded_by = data.get('recordedBy', '--')
    catalog_number = data.get('catalogNumber', '--')
    locality = data.get('locality', '--')
    dataset_name = data.get('datasetName', '--') # institudion ID ?
    return f'{locality}, {recorded_by} {record_number} ({dataset_name}:{catalog_number})'

def check_namespace_available(email, namespace_id):
    taicol_api = current_app.config['TAICOL_API']
    resp = requests.get(f'{taicol_api}/user/namespace?email={email}')
    if resp.ok:
        resp_json = resp.json()
        available_namespaces = resp_json.get('namespaces', [])
        if int(namespace_id) in available_namespaces:
            return True

    return False
